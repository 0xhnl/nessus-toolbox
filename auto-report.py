#!/usr/bin/env python3
"""Generate a docx vulnerability report from a Nessus HTML report.

Mimics the in-house docx template under ``doc-report/`` — Montserrat Semi Bold
headings, Lato body, Share Tech Mono code blocks. Groups findings into
Critical / High / Medium / Low Risk Findings.

Usage:
    python3 auto-report.py -f nessus-report.html -o output.docx
"""

import argparse
import re
import sys
from pathlib import Path

from bs4 import BeautifulSoup
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Bullet numbering from the template (numId=10 → abstractNumId=0 → Symbol "•").
BULLET_NUM_ID = 10

DEFAULT_TEMPLATE = (
    Path(__file__).resolve().parent
    / "doc-report"
    / "rpt.ysh-internal-assets-va-scan.20260610.docx"
)

# Severity bucket ordering for output.
SEVERITY_ORDER = ["Critical", "High", "Medium", "Low"]

# Primary signal — the plugin's "Risk Factor" field. Most reliable.
SEVERITY_BY_RISK_FACTOR = {
    "critical": "Critical",
    "high": "High",
    "medium": "Medium",
    "low": "Low",
}

# Fallback — header background color. Used only when Risk Factor is missing.
SEVERITY_BY_COLOR = {
    "#91243E": "Critical",
    "#DD4B50": "High",
    "#F18C43": "Medium",
    "#F8C851": "Low",
}

TOGGLE_RE = re.compile(r"toggleSection\('(id\d+-container)'\)")
HEADER_RE = re.compile(r"^\s*(\d+)\s*\((\d+)\)\s*-\s*(.+?)\s*-?\s*$")
CVE_RE = re.compile(r"CVE-\d{4}-\d{4,7}", re.IGNORECASE)


def severity_of(header, container):
    rf = section_text(container, "Risk Factor").strip().lower()
    sev = SEVERITY_BY_RISK_FACTOR.get(rf)
    if sev:
        return sev
    style = header.get("style") or ""
    m = re.search(r"background:\s*(#[0-9A-Fa-f]{6})", style)
    if m:
        return SEVERITY_BY_COLOR.get(m.group(1).upper())
    return None


def section_text(container, label):
    """Return the text of the div directly following a details-header with `label`."""
    for hdr in container.find_all("div", class_="details-header"):
        if hdr.get_text(strip=True) == label:
            sib = hdr.find_next_sibling()
            return sib.get_text("\n", strip=False) if sib else ""
    return ""


def host_blocks(container):
    """Yield (host_label, plugin_output_text) for each <h2> host header in container."""
    for h2 in container.find_all("h2"):
        label = h2.get_text(strip=True)
        parts = []
        sib = h2.next_sibling
        while sib is not None:
            if getattr(sib, "name", None) == "h2":
                break
            if getattr(sib, "name", None):
                cls = sib.get("class") or []
                if "clear" in cls:
                    sib = sib.next_sibling
                    continue
                txt = sib.get_text("\n", strip=False)
                if txt.strip():
                    parts.append(txt)
            sib = sib.next_sibling
        yield label, "\n".join(parts).rstrip()


def sort_cves(cves):
    def key(c):
        _, year, num = c.split("-")
        return (int(year), int(num))
    return sorted({c.upper() for c in cves}, key=key)


def split_paragraphs(text):
    """Split a description blob into clean paragraphs (collapse single newlines)."""
    if not text.strip():
        return []
    paragraphs = re.split(r"\n\s*\n", text)
    out = []
    for p in paragraphs:
        clean = re.sub(r"[ \t]+", " ", p).strip()
        # Preserve bullet-style lines (- foo, * foo, • foo) on their own lines.
        lines = [ln.strip() for ln in clean.split("\n") if ln.strip()]
        bulletish = sum(1 for ln in lines if re.match(r"^[-*•]\s+", ln))
        if lines and bulletish == len(lines):
            out.extend(lines)
        else:
            out.append(" ".join(lines))
    return out


# Imperative verbs → gerund. Covers the ~15 verbs Nessus uses to open its Solution text.
GERUND = {
    "upgrade": "upgrading",
    "update": "updating",
    "install": "installing",
    "disable": "disabling",
    "enable": "enabling",
    "reconfigure": "reconfiguring",
    "configure": "configuring",
    "apply": "applying",
    "remove": "removing",
    "replace": "replacing",
    "use": "using",
    "refer": "referring to",
    "restrict": "restricting",
    "ensure": "ensuring",
    "purchase": "purchasing",
    "contact": "contacting",
    "switch": "switching",
    "limit": "limiting",
    "set": "setting",
    "filter": "filtering",
}


def make_recommendation(solution):
    if not solution.strip():
        return "N/A"
    s = solution.strip()
    m = re.match(r"^(\w+)\b\s*(.*)$", s, re.DOTALL)
    if m:
        verb = m.group(1).lower()
        rest = m.group(2).strip()
        if verb in GERUND:
            tail = " " + rest if rest else ""
            return f"We recommend {GERUND[verb]}{tail}"
    return "We recommend " + s[0].lower() + s[1:]


def parse_report(html_path):
    with open(html_path, "rb") as f:
        soup = BeautifulSoup(f, "lxml")

    buckets = {s: [] for s in SEVERITY_ORDER}
    seen_titles = {s: set() for s in SEVERITY_ORDER}

    for header in soup.find_all("div", onclick=TOGGLE_RE):
        m = TOGGLE_RE.search(header.get("onclick", ""))
        if not m:
            continue
        container = soup.find("div", id=m.group(1))
        if container is None:
            continue

        sev = severity_of(header, container)
        if sev not in buckets:
            continue  # skip Info / unknown

        header_text = header.get_text(" ", strip=True)
        mh = HEADER_RE.match(header_text)
        title = mh.group(3).strip() if mh else header_text

        # Skip duplicates within the same severity bucket (same plugin merged twice).
        if title in seen_titles[sev]:
            continue
        seen_titles[sev].add(title)

        synopsis = section_text(container, "Synopsis").strip()
        description = section_text(container, "Description").strip()
        # Mirror the in-house template: synopsis sentence prepends the description.
        if synopsis and synopsis not in description:
            description = synopsis + "\n\n" + description
        solution = section_text(container, "Solution").strip()
        refs = " ".join([
            section_text(container, "References"),
            section_text(container, "See Also"),
        ])
        cves = sort_cves(CVE_RE.findall(refs))

        hosts = []
        outputs = []
        for host_label, output in host_blocks(container):
            hosts.append(host_label)
            if output:
                outputs.append(output)

        # Use the first non-empty per-host output as a representative sample.
        plugin_output = outputs[0] if outputs else ""

        buckets[sev].append({
            "title": title,
            "description": description,
            "plugin_output": plugin_output,
            "hosts": hosts,
            "solution": solution,
            "cves": cves,
        })

    return buckets


def strip_body(doc):
    body = doc.element.body
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def apply_bullet(paragraph, num_id=BULLET_NUM_ID, ilvl=0):
    """Attach numbering (bullet) to a paragraph via its w:numPr property."""
    pPr = paragraph._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl_el = OxmlElement("w:ilvl")
    ilvl_el.set(qn("w:val"), str(ilvl))
    numId_el = OxmlElement("w:numId")
    numId_el.set(qn("w:val"), str(num_id))
    numPr.append(ilvl_el)
    numPr.append(numId_el)
    pPr.append(numPr)


def add_lines(doc, text, style):
    # Strip leading/trailing blank lines; preserve interior ones.
    lines = [ln.rstrip() for ln in text.split("\n")]
    while lines and not lines[0].strip():
        lines.pop(0)
    while lines and not lines[-1].strip():
        lines.pop()
    for line in lines:
        doc.add_paragraph(line, style=style)


def build_doc(buckets, template_path, out_path):
    doc = Document(str(template_path))
    strip_body(doc)

    doc.add_paragraph("Findings", style="Heading 1")

    for sev in SEVERITY_ORDER:
        doc.add_paragraph(f"{sev} Risk Findings", style="Heading 2")
        for f in buckets[sev]:
            doc.add_paragraph(f["title"], style="Heading 3")

            doc.add_paragraph("Description", style="Heading 4")
            paragraphs = split_paragraphs(f["description"])
            if not paragraphs:
                paragraphs = ["N/A"]
            for para in paragraphs:
                doc.add_paragraph(para, style="Normal")

            if f["plugin_output"]:
                add_lines(doc, f["plugin_output"], "k-code-block")

            doc.add_paragraph("Affected Assets", style="Heading 4")
            for host in f["hosts"]:
                doc.add_paragraph(host, style="k-code-block")

            doc.add_paragraph("Recommendation", style="Heading 4")
            doc.add_paragraph(make_recommendation(f["solution"]), style="Normal")

            doc.add_paragraph("Vulnerability Reference", style="Heading 4")
            if f["cves"]:
                for cve in f["cves"]:
                    apply_bullet(doc.add_paragraph(cve, style="List Paragraph"))
            else:
                apply_bullet(doc.add_paragraph("N/A", style="List Paragraph"))

    doc.save(str(out_path))


def main():
    p = argparse.ArgumentParser(
        description="Generate a docx report from a Nessus HTML report."
    )
    p.add_argument("-f", "--file", required=True, help="Input Nessus HTML report")
    p.add_argument("-o", "--output", required=True, help="Output .docx file")
    p.add_argument("--template", default=str(DEFAULT_TEMPLATE),
                   help="Template docx used for fonts/styles "
                        f"(default: {DEFAULT_TEMPLATE.name})")
    args = p.parse_args()

    html_path = Path(args.file)
    if not html_path.is_file():
        sys.exit(f"Input not found: {html_path}")

    template_path = Path(args.template)
    if not template_path.is_file():
        sys.exit(f"Template not found: {template_path}")

    print(f"[+] Parsing {html_path}", file=sys.stderr)
    buckets = parse_report(html_path)

    total = 0
    for sev in SEVERITY_ORDER:
        n = len(buckets[sev])
        total += n
        print(f"    {sev:>8}: {n} finding(s)", file=sys.stderr)
    print(f"    {'Total':>8}: {total} finding(s)", file=sys.stderr)

    out_path = Path(args.output)
    build_doc(buckets, template_path, out_path)
    print(f"[+] Wrote {out_path}", file=sys.stderr)


if __name__ == "__main__":
    main()
